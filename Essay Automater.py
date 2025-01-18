import os
from collections import Counter
from tkinter import Tk, filedialog, Label, Button, Text, StringVar, Frame, OptionMenu, Radiobutton 
from tkinter.ttk import Progressbar
import tkinter.font as tkfont
import tkinter.messagebox as mbox
from tkinter.ttk import Style
from docx import Document
import re
from docx.shared import RGBColor

# Initialize variables
file_paths = []
get_if_keyword_is_prefix = 1

def select_files():
    global file_paths
    file_paths = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx")])
    if file_paths:
        selected_files_label.config(text=f"{len(file_paths)} file(s) selected",fg="green")
        update_preview()
    else:
        selected_files_label.config(text="No files selected")

def process_files():
    if not file_paths:
        selected_files_label.config(text="No files selected. Please select files first.", fg="red")
        return


    # Get user inputs
    text_to_append = text_entry.get("1.0", "end-1c").strip()
    prefix_pattern = prefix.get("1.0", "end-1c").strip()
    suffix_pattern = suffix.get("1.0", "end-1c").strip()
    if prefix_pattern: prefix_pattern = prefix_pattern + "_"
    if suffix_pattern: suffix_pattern = "_" + suffix_pattern 
    keyword_filename = [value.strip() for value in keyword_filename_text.get("1.0", "end-1c").strip().split(",")]


    keywords1 = [keyword.strip() + " " for keyword in keywords_1.get("1.0", "end-1c").splitlines() if keyword.strip()]
    keywords2 = [keyword.strip() + " " for keyword in keywords_2.get("1.0", "end-1c").splitlines() if keyword.strip()]
    keywords3 = [keyword.strip() + " " for keyword in keywords_3.get("1.0", "end-1c").splitlines() if keyword.strip()]

    keywords = keywords1 + keywords2 + keywords3

    if not prefix_pattern and not suffix_pattern and not keyword_filename:
        selected_files_label.config(text="Please provide a name pattern for renaming files.", fg="red")
        return
    
    #Check for overwrite
    overwriteFiles = ""
    for file_path in file_paths:
        keyword_prefix = ""
        keyword_suffix = ""
        dir_name, original_name = os.path.split(file_path)
        if keyword_filename != ['']: 
            freq_keyword = get_most_frequent_keyword(file_path) 
            if freq_keyword: #if a frequent keyword is found
                if get_if_keyword_is_prefix:
                    keyword_prefix = freq_keyword + "_"
                else:
                    keyword_suffix = "_" + freq_keyword
        new_name = f"{prefix_pattern}{keyword_prefix}{os.path.splitext(original_name)[0]}{keyword_suffix}{suffix_pattern}.docx"
        new_path = os.path.join(dir_name, new_name)
        if os.path.exists(new_path):
            overwriteFiles = overwriteFiles + new_path + "\n\n"
    if overwriteFiles:
        if not mbox.askokcancel("Overwrite Warning", "The following files will be overwritten:\n\n" + overwriteFiles):
            return


    # Define highlight color mapping
    color_mapping = {
        "Red": RGBColor(255, 0, 0),
        "Green": RGBColor(0, 255, 0),
        "Blue": RGBColor(0, 0, 255),
        "Yellow": RGBColor(255, 255, 0)
    }
    highlight_color1 = color_mapping.get(color_var1.get(), RGBColor(0, 0, 0))
    highlight_color2 = color_mapping.get(color_var2.get(), RGBColor(0, 0, 0))
    highlight_color3 = color_mapping.get(color_var3.get(), RGBColor(0, 0, 0))


    #start the progress bar
    progress.pack(padx=50, pady=30, side="top", fill="x")
    progress["value"] = 0
    progress["maximum"] = len(file_paths)
    root.update_idletasks()
    for file_path in file_paths:
        try:

            # Load the document
            doc = Document(file_path)

            # Highlight keywords in each paragraph
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                new_paragraph_text = original_text

                # Highlight all keywords by marking their positions and colors
                for keyword in keywords:
                    highlighted_keyword = f"[HIGHLIGHT_START_{keyword}]{keyword}[HIGHLIGHT_END]"
                    new_paragraph_text = re.sub(
                        re.escape(keyword), 
                        highlighted_keyword, 
                        new_paragraph_text, 
                        flags=re.IGNORECASE
                    )
                    
                    new_paragraph_text.replace(keyword, highlighted_keyword)

                # Clear all existing runs in the paragraph
                for run in paragraph.runs:
                    r = run._element
                    r.getparent().remove(r)

                # Rebuild the paragraph with the processed text
                sections = new_paragraph_text.split("[HIGHLIGHT_START_")
                for section in sections:
                    if "[HIGHLIGHT_END]" in section:
                        # Highlighted keyword
                        keyword_text_with_marker, remainder = section.split("[HIGHLIGHT_END]", 1)
                        keyword_text = keyword_text_with_marker.split("]", 1)[1]

                        highlighted_run = paragraph.add_run(keyword_text)
                        if keyword_text in keywords1:
                            highlighted_run.font.color.rgb = highlight_color1
                        elif keyword_text in keywords2:
                            highlighted_run.font.color.rgb = highlight_color2
                        elif keyword_text in keywords3:
                            highlighted_run.font.color.rgb = highlight_color3
                        highlighted_run.font.bold = True

                        # Add the remainder (non-highlighted text)
                        if remainder:
                            paragraph.add_run(remainder)
                    else:
                        # Normal text section
                        paragraph.add_run(section)


            #Add boiler plate
            doc.add_paragraph(text_to_append)

            # Save the document with the new name
            keyword_prefix = ""
            keyword_suffix = ""
            if keyword_filename != ['']: 
                freq_keyword = get_most_frequent_keyword(file_path) 
                if freq_keyword: #if a frequent keyword is found
                    if get_if_keyword_is_prefix:
                        keyword_prefix = freq_keyword + "_"
                    else:
                        keyword_suffix = "_" + freq_keyword
           
            dir_name, original_name = os.path.split(file_path)
            new_name = f"{prefix_pattern}{keyword_prefix}{os.path.splitext(original_name)[0]}{keyword_suffix}{suffix_pattern}.docx"
            new_path = os.path.join(dir_name, new_name)
            doc.save(new_path)
            progress["value"] += 1
            root.update_idletasks()
            print(f"Updated and saved: {new_path}")
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
    
    selected_files_label.config(text="All files created!", fg="Green")

def update_preview():
    keyword_is_prefix = get_if_keyword_is_prefix
    keyword_filename = [value.strip() for value in keyword_filename_text.get("1.0", "end-1c").strip().split(",")]
    prefix_pattern = prefix.get("1.0", "end-1c").strip()
    suffix_pattern = suffix.get("1.0", "end-1c").strip()
    keyword_prefix = ""
    keyword_suffix = ""
    if prefix_pattern: prefix_pattern = prefix_pattern + "_"
    if suffix_pattern: suffix_pattern = "_" + suffix_pattern 
    if file_paths:
        preview_text.delete("1.0", "end")
        for file_path in file_paths:
            original_name = os.path.split(file_path)[1]
            if keyword_filename != ['']: 
                freq_keyword = get_most_frequent_keyword(file_path) 
                if freq_keyword: #if a frequent keyword is found
                    if keyword_is_prefix:
                        keyword_prefix = freq_keyword + "_"
                    else:
                        keyword_suffix = "_" + freq_keyword
            new_name = f"{prefix_pattern}{keyword_prefix}{os.path.splitext(original_name)[0]}{keyword_suffix}{suffix_pattern}.docx"
            preview_text.insert("end", new_name + "\n")

def get_most_frequent_keyword(file):
    keyword_filename = [value.strip() for value in keyword_filename_text.get("1.0", "end-1c").strip().split(",")]

    if keyword_filename != []:    
        # Load the .docx file
        document = Document(file)
        # Read all text from the document
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + " "
        
        # Normalize the text and count occurrences of keywords
        text = text.lower()
        keyword_filename = [keyword.lower() for keyword in keyword_filename]
        # keyword_counts = Counter(word for word in text.split() if word in keyword_filename)
        keyword_counts = Counter()
        for phrase in keyword_filename:
            count = text.count(phrase)  # Count occurrences of the entire phrase
            if count > 0:
                keyword_counts[phrase] = count
        # Return the most common keyword, or None if no matches
        if keyword_counts:
            return keyword_counts.most_common(1)[0][0].capitalize()
        return None


def toggle_theme():
    global dark_mode
    dark_mode = not dark_mode
    bg_color, fg_color, active_bg = ("#2E2E2E", "#FFFFFF", "#444444") if dark_mode else ("#AAAAAA", "#000000", "#DDDDDD")

    root.tk_setPalette(background=bg_color, foreground=fg_color, activeBackground=active_bg, activeForeground=fg_color)
    style.configure("TButton", background=active_bg, foreground=fg_color)
    style.configure("TLabel", background=bg_color, foreground=fg_color)
    style.configure("TEntry", background=active_bg, foreground=fg_color, insertcolor=fg_color)

    # Update text areas
    text_entry.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    preview_text.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    prefix.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    suffix.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    keywords_1.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    keywords_2.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    keywords_3.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    keyword_filename_text.config(background=active_bg, foreground=fg_color, insertbackground=fg_color)
    
def change_get_if_keyword_is_prefix(myBool):
    global get_if_keyword_is_prefix
    get_if_keyword_is_prefix = myBool
    update_preview()


# Initialize the GUI
root = Tk()
root.title("Word Document Batch Updater")
root.geometry("1280x1024")

# Apply dark mode style
style = Style(root)
dark_mode = False

# Define a custom font
default_font = tkfont.nametofont("TkDefaultFont")
default_font.configure(size=14)  # Change the font size here
root.option_add("*Font", default_font)

# Create UI components
top_frame = Frame(root)
top_frame.pack(fill="x", side= "top")

Button(top_frame, text="Toggle Darkmode", command=toggle_theme, padx=10, pady=5).pack(side="left", padx=10, pady=10)
Button(top_frame, text="Process Files", command=process_files).pack(side="right", padx=10, pady=10)
progress = Progressbar(top_frame, orient="horizontal", mode="determinate")
progress.pack(padx=50, pady=30, side="top", fill="x")
progress.pack_forget()


selected_files_label = Label(root, text="No files selected", fg="red")
selected_files_label.pack(pady=5)

Button(root, text="Select Files", command=select_files ,padx=10, pady=5).pack(pady=5)

Label(root, text="Enter the text to append inside file:").pack(pady=5)
text_entry = Text(root, wrap="word", height=5)
text_entry.pack(fill="both", expand=True, pady=5, padx=20)

keyword_name_frame = Frame(root)
keyword_name_frame.pack(fill="x", side="top", pady=5, padx=20) 
Label(keyword_name_frame, text="Name files for the most frequent keyword:\n(Comma seperated)", justify="center").pack(pady=5, side="left")

radio_frame = Frame(keyword_name_frame)
radio_frame.pack(fill="x", side="right") 

r1 = Radiobutton(radio_frame, text="Prefix", value=1, variable=get_if_keyword_is_prefix, command=lambda: change_get_if_keyword_is_prefix(1), indicatoron=0,selectcolor="#2e2e2e", padx=5, pady=5)
r1.pack(pady=5, padx=5, side="top")
r1.invoke()
r2 = Radiobutton(radio_frame, text="Suffix", value=0, variable=get_if_keyword_is_prefix, command=lambda: change_get_if_keyword_is_prefix(0), indicatoron=0,selectcolor="#2e2e2e", padx=5, pady=5)
r2.pack(pady=5, padx=5, side="bottom")

keyword_filename_text = Text(keyword_name_frame, width=30, height=1)
keyword_filename_text.pack(padx=20, pady=40, fill="both")
keyword_filename_text.bind("<KeyRelease>", lambda e: update_preview())


name_frames = Frame(root)
name_frames.pack(fill="x", side="top", pady=5, padx=20) 

pre_frame = Frame(name_frames)
pre_frame.pack(side="left",padx=10, pady=10) 
Label(pre_frame, text="Enter the new file prefix:").pack(pady=5, side="top")
prefix = Text(pre_frame, width=30, height=1)
prefix.pack(pady=5, side="bottom")
prefix.bind("<KeyRelease>", lambda e: update_preview())

post_frame = Frame(name_frames)
post_frame.pack(side="right",padx=10, pady=10) 
Label(post_frame, text="Enter the new file suffix:").pack(pady=5, side="top")
suffix = Text(post_frame, width=30, height=1)
suffix.pack(pady=5, side="bottom")
suffix.bind("<KeyRelease>", lambda e: update_preview())

newfile_frame = Frame(name_frames)
newfile_frame.pack(padx=10, pady=10, fill="both") 
Label(newfile_frame, text="Preview of new filenames:").pack(padx=5, pady=5,side="top", fill="both")
preview_text = Text(newfile_frame, state="normal", height=6)
preview_text.pack(pady=5, side="bottom", fill="both")




highlight_frame = Frame(root)
highlight_frame.pack(side="bottom",padx=10, pady=10) 
Label(highlight_frame, text="Enter keywords to highlight (separated by newlines):").pack(pady=5)

keyword_frame_1 = Frame(highlight_frame)
keyword_frame_1.pack(padx=10, pady=10, side="left") 
Label(keyword_frame_1, text="Select highlight color 1:").pack(pady=5, side="top")
color_var1 = StringVar(value="Red")
color_dropdown = OptionMenu(keyword_frame_1, color_var1, "Red", "Green", "Blue", "Yellow")
color_dropdown.pack(pady=5)
keywords_1 = Text(keyword_frame_1, height=5, width=30)
keywords_1.pack(pady=5, side="bottom", fill="x", expand=True)


keyword_frame_3 = Frame(highlight_frame)
keyword_frame_3.pack(padx=10, pady=10, side="right") 
Label(keyword_frame_3, text="Select highlight color 3:").pack(pady=5, side="top")
color_var2 = StringVar(value="Green")
color_dropdown = OptionMenu(keyword_frame_3, color_var2, "Red", "Green", "Blue", "Yellow")
color_dropdown.pack(pady=5)
keywords_3 = Text(keyword_frame_3, wrap="word", height=5, width=30)
keywords_3.pack(pady=5, side="bottom", fill="x", expand=True)

keyword_frame_2 = Frame(highlight_frame)
keyword_frame_2.pack(padx=10, pady=10) 
Label(keyword_frame_2, text="Select highlight color 2:").pack(pady=5, side="top")
color_var3 = StringVar(value="Blue")
color_dropdown = OptionMenu(keyword_frame_2, color_var3, "Red", "Green", "Blue", "Yellow")
color_dropdown.pack(pady=5)
keywords_2 = Text(keyword_frame_2, wrap="word", height=5, width=30)
keywords_2.pack(pady=5, side="bottom", fill="x", expand=True)




toggle_theme()
# Run the GUI event loop
root.mainloop()
